#!/usr/bin/env python3
# -*- coding: utf-8 -*-
from docx import Document
import os

base_path = os.path.dirname(__file__)

# Arquivo 1
print("="*80)
print("ARQUIVO 1: ForcaTotal_Historia_Operabilidade.docx")
print("="*80)
doc1 = Document(os.path.join(base_path, 'docs/ForcaTotal_Historia_Operabilidade.docx'))
texto1 = '\n'.join([p.text for p in doc1.paragraphs if p.text.strip()])
print(texto1)

# Arquivo 2
print("\n" + "="*80)
print("ARQUIVO 2: ForcaTotal_Objetivos_Requisitos_Iniciais.docx")
print("="*80)
doc2 = Document(os.path.join(base_path, 'docs/ForcaTotal_Objetivos_Requisitos_Iniciais.docx'))
texto2 = '\n'.join([p.text for p in doc2.paragraphs if p.text.strip()])
print(texto2)
